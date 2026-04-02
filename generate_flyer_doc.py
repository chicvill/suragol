# -*- coding: utf-8 -*-
import os
import sys
import subprocess

def install_and_import():
    try:
        import docx
    except ImportError:
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
        import docx
    return docx

docx = install_and_import()
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = docx.Document()

# 문서 제목
title = doc.add_heading('식당 운영의 혁신, MQnet 스마트 QR 대기/주문 시스템', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph('비싼 하드웨어 NO! 우리 매장에 딱 맞는 가장 스마트한 통제 솔루션\n')
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 1. 사장님의 고민
doc.add_heading('1. 바쁜 피크타임, 이런 고민 없으신가요?', level=1)
pains = [
    "점심/저녁 시간만 되면 웨이팅 손님 챙기랴, 홀 서빙하랴 직원들이 쉴 틈이 없습니다.",
    "대기 손님들이 밖에서 마냥 기다리다 지쳐 그냥 돌아가시는(노쇼) 경우가 생깁니다.",
    "대기 명단을 종이에 적고 육성으로 부르다 보니 누락이나 컴플레인이 발생합니다.",
    "테이블마다 주문용 패드를 놓자니 수백만 원의 초기 기기값과 향후 고장이 너무 걱정됩니다."
]
for p_text in pains:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(p_text)

# 2. MQnet이 제안하는 완벽한 해결책
doc.add_heading('2. MQnet 시스템 특장점 (왜 선택해야 할까요?)', level=1)
adv = [
    ("초기 기기 다량 구매 비용 0원", "고가의 키오스크나 전용 태블릿을 살 필요 없이, 테이블과 입구에 예쁜 QR 스티커만 붙여두세요. 손님이 들고 온 스마트폰이 곧 고성능 오더기이자 진동벨이 됩니다."),
    ("앱 설치 없는 1초 자동 알림 (SMS)", "손님이 QR을 찍기만 하면 앱 설치 없이 바로 대기/주문 등록! 대기 순서가 다가오거나(예: 3팀 앞) 입장하실 때 자동으로 사장님 폰에서 손님 폰으로 문자가 발송되어 매장 앞 혼잡을 완전히 없앱니다."),
    ("포스기(POS) 화면과 실시간 100% 동기화", "사장님 카운터 화면을 새로고침 할 필요가 전혀 없습니다. 손님이 주문하거나 대기를 걸면 카운터 패드에 즉시 소리와 함께 명단이 떠서 즉각적인 대처가 가능합니다."),
    ("인건비 대폭 절감", "대기 명단을 펜으로 적고 밖으로 나가 손님을 부르는 서빙 알바 1명 이상의 역할을 톡톡히 해냅니다. 직원들은 음식 세팅과 친절한 서비스에만 온전히 집중할 수 있습니다.")
]
for title_text, desc in adv:
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(title_text + " : ")
    r.bold = True
    try:
        r.font.color.rgb = RGBColor(0, 51, 153) # 진한 파란색 강조
    except:
        pass
    p.add_run(desc)

# 3. 합리적인 도입 비용
doc.add_heading('3. 거품을 완전히 뺀 파격적인 비용', level=1)
cost_p = doc.add_paragraph('보통 매장 테이블 10개에 주문 패드를 설치하려면 수백만 원이 우습게 깨집니다. MQnet은 철저한 소프트웨어 고도화를 통해 시스템 도입의 벽을 확실히 낮췄습니다.\n')

cost_items = [
    ("■ 초기 세팅비 (최초 1회 한정)", "약 20만원 대\n(안내: QR 테이블 아크릴 인쇄 실비 + 메뉴 DB 등록 및 화면 디자인 구성 + 자동 문자 서버 초기 연동비 포함. ※ 매장 메뉴 사진을 원본으로 자체 제공해주시면 세팅비가 대폭 할인됩니다!)"),
    ("■ 월 운영 유지비", "월 3~5만 원 수준\n(안내: 클라우드 서버 유지, 고정 도메인 트래픽, 손님 자동 대기 안내 문자(SMS) 발생 실비가 전부 포함된 올인원 비용. 알바생 단 3~4시간 시급이면 매달 복잡한 홀 통제가 끝납니다.)")
]
for title_text, desc in cost_items:
    p = doc.add_paragraph()
    r = p.add_run(title_text + "\n")
    r.bold = True
    p.add_run("   " + desc + "\n")

# 4. 마무리 및 시연
doc.add_heading('4. 1일 무료 시연 시스템 체험', level=1)
doc.add_paragraph('백문이 불여일견입니다! 지금 바로, 사장님의 스마트폰 카메라를 켜서 1분 만에 작동 화면을 경험해보세요.')
contact = doc.add_paragraph()
cr = contact.add_run('\n▶ 담당 매니저 (MQnet) : O O O\n▶ 상담 및 연락처 : 010-XXXX-XXXX\n▶ 이메일 : \n')
cr.bold = True

footer = doc.add_paragraph('\n※ "사장님은 요리에만 집중하십시오. 골치 아픈 홀 관리는 MQnet이 책임지겠습니다!"')
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in footer.runs:
    r.bold = True
    try:
        r.font.color.rgb = RGBColor(204, 0, 0)
    except:
        pass

file_name = 'MQnet_식당배포용_제안서.docx'
file_path = os.path.join(r'c:\Users\USER\Dev\왕궁중화요리', file_name)
try:
    doc.save(file_path)
    print(f"===========================================================")
    print(f"✨ 사장님께 배포할 [MQnet_식당배포용_제안서.docx] 문서가 성공적으로 생성되었습니다!")
    print(f"👉 파일 위치: {file_path}")
    print(f"===========================================================")
except Exception as e:
    print(f"파일을 생성할 수 없습니다: {e}")
