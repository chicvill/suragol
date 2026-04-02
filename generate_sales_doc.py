# -*- coding: utf-8 -*-
import os
import sys
import subprocess

def install_and_import():
    try:
        import docx
    except ImportError:
        print("python-docx 패키지가 없어서 자동으로 설치합니다...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
        import docx
    return docx

docx = install_and_import()
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = docx.Document()

# 문서 제목
title = doc.add_heading('MQnet 스마트 웨이팅/주문 시스템 영업 가이드', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('영업사원 배포용 및 식당 사장님 제안 실전 가이드라인\n')

# 1. 핵심 세일즈 포인트
doc.add_heading('1. 핵심 세일즈 포인트 (사장님 제안 소구점)', level=1)
points = [
    ("초기 기기 도입 비용 & 유지보수 제로 (No Kiosk, No Pad)", 
     "비싼 키오스크 기계나 대기 명단용 태블릿을 수십만 원 주고 구매하실 필요가 없습니다. 고객의 스마트폰이 곧 고성능 테이블 오더기이자 진동벨이 됩니다."),
    ("앱 설치 없는 직관적인 자동 문자(SMS) 알림", 
     "복잡한 앱 설치나 회원가입 시킬 필요 없이 손님이 QR만 찍으면 끝입니다. 내 앞 대기팀이 3팀 남았을 때, 그리고 입장할 때 자동으로 문자가 발송되어 손님들이 매장 앞에서 서성이지 않아도 됩니다."),
    ("홀 인건비 절감 및 매장 혼잡도 완화", 
     "바쁜 피크 타임에 직원들이 육성으로 손님 부르고, 명단 적고 확인하러 다니는 시간을 완전히 없애줍니다. 직원 1명 몫을 시스템이 대신해 줍니다."),
    ("카운터 100% 실시간 통제 (새로고침 불필요)", 
     "사장님 카운터 포스기나 패드 화면에 대기 현황이 실시간으로 동기화됩니다. 새로고침 할 필요 없이 손님이 밖에서 QR을 찍는 즉시 카운터 화면에 '띠링' 하고 명단이 올라옵니다.")
]
for title_text, desc in points:
    p = doc.add_paragraph(style='List Bullet')
    runner = p.add_run(title_text)
    runner.bold = True
    p.add_run(f'\n{desc}')

# 2. 실전 영업 스크립트
doc.add_heading('2. 실전 영업 스크립트 (상황별 예시)', level=1)
scripts = [
    ("1단계: 첫 인사 및 훅 (사장님 고충 공감하기)",
     '"안녕하세요 사장님, MQnet에서 나왔습니다! 요즘 피크시간에 손님들 몰릴 때 홀 관리하시기 참 힘드시죠? 특히 대기 줄 길어지면 손님들도 불만이 생기고, 직원분들이 명단 관리하느라 너무 정신이 없잖아요. 혹시 사장님 매장도 이런 고민 있으신가요?"'),
    ("2단계: 솔루션 제시 (쉽고 명확하게)",
     '"그래서 저희 MQnet에서 단 1초 만에 스마트폰으로 끝나는 \'QR 웨이팅 및 주문 시스템\'을 도입해 드리고 있습니다. 테이블이나 입구에 예쁜 QR코드 스티커만 붙여두시면 끝입니다."'),
    ("3단계: 디테일 및 혜택 설명 (시스템 강점 어필)",
     '"등록만 하면 차례가 다가올 때 자동으로 문자가 가기 때문에 노쇼(No-show)도 획기적으로 줄어듭니다. 또한 카운터 화면을 통해 [입장 확인] 버튼 한 번만 누르시면 관리가 끝납니다."'),
    ("4단계: 클로징 및 시연 유도",
     '"지금 이걸 도입하신 매장들은 홀 서빙 직원 1명 더 쓴 것 같다고 만족하십니다. 사장님, 말로 설명 듣는 것보다 제가 사장님 스마트폰으로 지금 바로 1분 만에 보여드릴까요?"')
]
for step_title, script_content in scripts:
    doc.add_heading(step_title, level=2)
    p = doc.add_paragraph()
    runner = p.add_run(script_content)
    runner.italic = True
    try:
        runner.font.color.rgb = RGBColor(0, 102, 204)
    except:
        pass

# 3. 예상 소요 장비 및 도입 비용 (최소 비용 산출)
doc.add_heading('3. 예상 소요 장비 및 도입 비용 (최소 비용 산출 기준)', level=1)
doc.add_paragraph('비싼 하드웨어(키오스크/테이블패드) 대신 소프트웨어 최적화로 초기 도입 장벽을 극적으로 낮추었습니다. 영업 시 아래 가이드라인을 기준으로 사장님들께 제안하시길 바랍니다.')

doc.add_heading('■ 소요 장비 (하드웨어)', level=2)
equip_points = [
    ("매장 카운터용 통제 PC 또는 태블릿", "기존 매장에 있는 포스(POS)기 또는 사장님/직원용 스마트폰, 패드 그대로 활용 가능 (장비 구매비 0원)"),
    ("QR 코드 테이블 스티커 및 입구용 아크릴 보드", "최소 인쇄비 매장당 약 3~5만 원 소요 (디자인 인쇄 및 아크릴 스탠드 제작비)"),
    ("매장 내 무선 인터넷 환경", "매장 공유기 활용 (대부분 구축되어 있음)")
]
for title_text, desc in equip_points:
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(title_text + " : ")
    r.bold = True
    p.add_run(desc)

doc.add_heading('■ 시스템 초기 세팅 및 개발비 (1회성)', level=2)
dev_points = [
    ("식당 맞춤형 메뉴판 세팅 (DB 등록 및 UI 구성)", "약 10 ~ 15만 원 (메뉴 가짓수 및 옵션 복잡도에 따라 상이)"),
    ("메뉴 및 매장 전경 사진 전문 촬영비", "약 10 ~ 20만 원 (최소 기본 컷 기준, ※ 사장님이 기존 사진 원본을 자체적으로 제공하시면 무상 인하 가능)"),
    ("서버 안정화 및 문자(SMS)/알림톡 연동 셋업", "약 10만 원 (도메인 연결, Cloudflare 보안 터널링 및 알리고 API 초기 세팅)")
]
for title_text, desc in dev_points:
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(title_text + " : ")
    r.bold = True
    p.add_run(desc)

p_total = doc.add_paragraph()
p_total.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r_total = p_total.add_run('▶ 초기 세팅비 합계: 약 20만 원 ~ 45만 원 (식당 메뉴판 자체 제공 시 가장 저렴)')
r_total.bold = True
try:
    r_total.font.color.rgb = RGBColor(204, 0, 0)
except:
    pass

doc.add_heading('■ 시스템 유지보수 비용 (매월 과금 제안)', level=2)
p_monthly = doc.add_paragraph('클라우드 서버 유지, 고정 도메인, 테이블 실시간 웹소켓(Socket) 트래픽 및 자동 문자 발송 비용을 모두 포괄하여 ')
r_month = p_monthly.add_run('월 3만 원 ~ 5만 원')
r_month.bold = True
r_month.underline = True
p_monthly.add_run(' 수준의 구독형 관리 모델로 제안하십시오. 키오스크 렌탈비(보통 1대에 1.5만원~3만원) 대비 파격적인 가격이며, 대기줄 한 팀만 안 놓쳐도 월 이용료 뽑고 남는다고 어필하시는 것이 핵심입니다.')

# 4. 영업 현장 시연 팁
doc.add_heading('4. 영업 현장 시연(Demo) 팁', level=1)
demo_tips = [
    "사장님 스마트폰 카메라로 QR을 찍고 메뉴 선택/대기를 직접 체험시키기",
    "예시로 대기를 걸었을 때, 접수 문자가 거의 1초 만에 스마트폰으로 오는 것을 직접 확인시키기",
    "카운터용 화면(영업사원 패드)에 사장님이 입력한 내역이 새로고침 없이 즉시 연동되는 모습 눈으로 보여주기"
]
for idx, tip in enumerate(demo_tips, 1):
    doc.add_paragraph(f"{idx}. {tip}")

# 마무리
doc.add_paragraph('\n')
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
bold_footer = footer.add_run('※ 기기를 파는 것이 아니라, "식당의 평화"와 "인건비 절감"을 파는 것임을 명심하세요. 영업 파이팅!')
bold_footer.bold = True
try:
    bold_footer.font.color.rgb = RGBColor(0, 51, 102)
except:
    pass

file_name = 'MQnet_QR시스템_영업가이드.docx'
file_path = os.path.join(r'c:\Users\USER\Dev\왕궁중화요리', file_name)
try:
    doc.save(file_path)
    print(f"===========================================================")
    print(f"✨ 비용 산출 파트가 추가된 MS Word 파일이 갱신되었습니다!")
    print(f"👉 파일 저장 완료: {file_path}")
    print(f"===========================================================")
except Exception as e:
    print(f"파일을 저장하려면 기존 Word 파일이 닫혀있어야 합니다: {e}")
